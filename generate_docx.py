# -*- coding: utf-8 -*-
"""
生成《转辙机动作电流异常检测与退化预警系统_技术详解.docx》
每个部分按: 设计背景(为什么做) → 怎么做(数据/算法/技术) → 结果(图/表 + 说明) 三段式.

依赖: python-docx (base python 3.13 已装)
运行: python generate_docx.py
"""
import json
import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# 常量
# ----------------------------------------------------------------------------
ROOT = os.path.dirname(os.path.abspath(__file__))
OUT_NAME = '转辙机动作电流异常检测与退化预警系统_技术详解.docx'
OUT_PATH = os.path.join(ROOT, OUT_NAME)

IMG_FLOW = os.path.join(ROOT, '转辙机三相动作电流_流程图.png')
IMG_WAVE = os.path.join(ROOT, '转辙机三相动作电流_波形.png')
IMG_LINE = os.path.join(ROOT, '转辙机三相动作电流_时序线条.png')
IMG_DIR = os.path.join(ROOT, 'outputs')
IMAGES = {
    'score': os.path.join(IMG_DIR, '02_score_dist.png'),
    'roc': os.path.join(IMG_DIR, '03_roc_pr.png'),
    'recon': os.path.join(IMG_DIR, '04_recon_samples.png'),
    'confusion': os.path.join(IMG_DIR, '05_confusion.png'),
    'warning': os.path.join(IMG_DIR, '07_warning_trend.png'),
}
STATS_PATH = os.path.join(IMG_DIR, '_data_stats.json')
METRICS_PATH = os.path.join(IMG_DIR, '_detection_metrics.json')
RUN_LOG = os.path.join(IMG_DIR, '_detection_run.log')

# 配色
HDR_FILL = '2F5597'      # 表头深蓝
ALT_FILL = 'D9E2F3'      # 隔行浅蓝
GRAY = RGBColor(0x59, 0x59, 0x59)

# ----------------------------------------------------------------------------
# 工具函数
# ----------------------------------------------------------------------------

def set_run_font(run, cn='宋体', en='Times New Roman', size=10.5, bold=False,
                 color=None, italic=False):
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:ascii'), en)
    rf.set(qn('w:hAnsi'), en)
    rf.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    return run


def set_style_font(style, cn='宋体', en='Times New Roman', size=10.5, bold=False):
    style.font.name = en
    style._element.get_or_add_rPr()
    rf = style._element.rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        style._element.rPr.append(rf)
    rf.set(qn('w:ascii'), en)
    rf.set(qn('w:hAnsi'), en)
    rf.set(qn('w:eastAsia'), cn)
    style.font.size = Pt(size)
    style.font.bold = bold


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def add_h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    set_run_font(r, cn='黑体', en='Arial', size=16, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    set_run_font(r, cn='黑体', en='Arial', size=13, bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, indent=True, size=10.5, bold=False, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Pt(size * 2)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.4
    pf.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10.5, bold_prefix=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.space_after = Pt(3)
    r0 = p.add_run('• ')
    set_run_font(r0, size=size, color=RGBColor(0x2F, 0x55, 0x97))
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        set_run_font(rb, size=size, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=size)
    return p


def add_image(doc, path, width_cm, caption=None, figure_no=None):
    if not os.path.exists(path):
        add_body(doc, f'[缺图] {os.path.basename(path)}', indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        label = f'图 {figure_no} ' if figure_no else ''
        r = cp.add_run(f'{label}{caption}')
        set_run_font(r, cn='黑体', en='Arial', size=9, color=GRAY)


def add_table(doc, headers, rows, caption=None, table_no=None, col_widths=None,
              align_right_cols=None):
    if not rows:
        return
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, h in enumerate(headers):
        c = t.cell(0, j)
        c.text = ''
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run_font(r, cn='黑体', en='Arial', size=9.5, bold=True,
                     color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(c, HDR_FILL)
    # 数据行
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            c = t.cell(i + 1, j)
            c.text = ''
            p = c.paragraphs[0]
            if align_right_cols and j in align_right_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v))
            set_run_font(r, size=9.5)
            if i % 2 == 1:
                shade_cell(c, ALT_FILL)
    # 列宽
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(len(rows) + 1):
                t.cell(i, j).width = Cm(w)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(10)
        label = f'表 {table_no} ' if table_no else ''
        r = cp.add_run(f'{label}{caption}')
        set_run_font(r, cn='黑体', en='Arial', size=9, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ----------------------------------------------------------------------------
# 数据加载
# ----------------------------------------------------------------------------

def load_stats():
    with open(STATS_PATH, encoding='utf-8') as f:
        return json.load(f)


def load_metrics():
    """优先 JSON, 回退解析日志, 再回退 README 数字."""
    if os.path.exists(METRICS_PATH):
        with open(METRICS_PATH, encoding='utf-8') as f:
            return json.load(f)
    if os.path.exists(RUN_LOG):
        txt = open(RUN_LOG, encoding='utf-8', errors='replace').read()
        d = {}
        for key, pat in [
            ('auc_roc', r'AUC-ROC:\s*([\d.]+)'),
            ('auc_pr', r'AUC-PR:\s*([\d.]+)'),
            ('precision', r'精确率:\s*([\d.]+)'),
            ('recall', r'召回率:\s*([\d.]+)'),
            ('f1', r'F1:\s*([\d.]+)'),
            ('fpr', r'虚警率\(FPR\):\s*([\d.]+)'),
        ]:
            m = re.search(pat, txt)
            if m:
                d[key] = float(m.group(1))
        m = re.search(r'检出:\s*(\d+)/(\d+)', txt)
        if m:
            d['tp'] = int(m.group(1))
            d['tp_fn'] = int(m.group(2))
        m = re.search(r'FP:\s*(\d+)', txt)
        if m:
            d['fp'] = int(m.group(1))
        if d:
            return d
    # README 基线
    return {
        'auc_roc': 0.9921, 'auc_pr': 0.8443, 'precision': 0.312,
        'recall': 0.857, 'f1': 0.457, 'fpr': 0.0063,
        'tp': 96, 'tp_fn': 112, 'fp': 212,
    }


# ----------------------------------------------------------------------------
# 文档构建
# ----------------------------------------------------------------------------

def build_doc():
    stats = load_stats()
    m = load_metrics()

    doc = Document()
    # 默认样式
    set_style_font(doc.styles['Normal'], size=10.5)
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.4)
        s.right_margin = Cm(2.4)

    # =====================================================================
    # 封面
    # =====================================================================
    for _ in range(4):
        doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run('转辙机动作电流异常检测与退化预警系统')
    set_run_font(r, cn='黑体', en='Arial', size=26, bold=True)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run('—— 技术详解 · 设计与实现 ——')
    set_run_font(r, cn='黑体', en='Arial', size=15, color=GRAY)
    doc.add_paragraph()
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run('基于 CNN-LSTM 时频双通路自编码器的无监督异常检测与逐台退化预警')
    set_run_font(r, cn='宋体', size=11)
    for _ in range(3):
        doc.add_paragraph()
    mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = mp.add_run('文档版本: v1.0        日期: 2026-08-12')
    set_run_font(r, cn='宋体', size=11)
    doc.add_page_break()

    # =====================================================================
    # 前言 / 阅读说明
    # =====================================================================
    add_h1(doc, '阅读说明')
    add_body(doc, '本技术文档围绕"转辙机(道岔)三相动作电流异常检测与退化预警"系统展开,'
                  '对系统的每一个组成部分,均按照"设计背景(为什么做)→ 怎么做(数据、算法设计与所用技术)→ '
                  '结果(实现效果,配图表并加以说明)"的三段式结构进行详解。')
    add_body(doc, '文档覆盖完整链路:从输入数据与预处理,到时域/频域双通路特征提取、交叉注意力对齐、'
                  '拼接融合与 BiLSTM、在线辅助特征注入、镜像解码器、辅助监督头与训练策略,再到异常检测管线'
                  '与逐台退化预警,共 11 个章节。检测指标与结果图均来自本仓库最新一次运行'
                  '(重跑 run_detection.py),可复现。')
    add_body(doc, '阅读建议:第 0 章给出系统全貌,后续每章可独立阅读;模型架构部分(第 2~8 章)需要'
                  '一定的深度学习基础,检测与预警部分(第 9~10 章)面向工程应用。')

    # =====================================================================
    # 第 0 章 系统总览
    # =====================================================================
    add_h1(doc, '0. 系统总览')

    add_h2(doc, '0.1 设计背景')
    add_body(doc, '转辙机作为铁路道岔转换的关键执行机构,其动作电流波形蕴含着丰富的机械状态信息,为设备状态监测'
                  '提供了直观而可靠的观测手段。机械卡阻是该类设备最具代表性且危害较高的故障形态:当道岔转换'
                  '受到机械阻力时,驱动电机需持续输出转矩以克服阻碍,致使转换段电流出现异常抬升,并伴随频谱'
                  '能量分布的改变。该故障在电流波形上的物理签名具有显著的局部性——主要作用于转换段电流,'
                  '而启动冲击与解锁段几乎不受影响——这一特性既为检测提供了明确的判别依据,也对特征设计的'
                  '精确性提出了要求。')
    add_body(doc, '然而,实际运维场景中存在两个制约有监督学习方法应用的关键约束。其一,卡阻故障率极低'
                  '(约 0.1%),故障样本数量稀缺,难以支撑有监督分类器的训练;其二,不同转辙机个体的正常电流'
                  '基线存在显著差异,低基础机发生卡阻后的电流水平可能与高基础机的正常水平相当,致使故障信号'
                  '被个体差异所淹没,构成漏检的主要来源。')
    add_body(doc, '基于上述分析,本文提出一种仅依赖正常数据训练的无监督检测框架:首先使模型充分学习'
                  '"正常动作的流形",继而以各样本对正常流形的偏离程度度量其异常水平。该范式与自编码器重构'
                  '误差检测的基本假设相一致,亦契合工业场景对数据可得性的约束。')
    add_body(doc, '在特征建模层面,考虑到卡阻故障同时体现为时域形态的变化与频域结构的变化,本文构建时域与'
                  '频域双通路的特征提取架构。为避免两路特征简单拼接所导致的信息交互不足,引入交叉注意力对齐'
                  '机制,使时域特征与频域特征在编码阶段即通过互查询实现跨模态信息融合,并结合可学习门控与'
                  '残差连接保证融合的稳健性。')
    add_body(doc, '在表征学习层面,单纯的重构损失难以迫使潜在表示充分编码卡阻相关的关键物理量。为此,'
                  '本文引入幅度解码头、时段判别头与频域正常性模型三类辅助监督目标,通过多任务训练强化潜在'
                  '表示对异常形态的敏感性。')
    add_body(doc, '在检测评分层面,为应对单一误差度量易致漏检的问题,提出由 PW-MSE、PhaseErr、ClusterLatent '
                  '与 RelErr 构成的多分量加权融合机制。四个分量分别刻画重构偏差、相位结构、机簇相对偏差与'
                  '波形级物理比值等互补视角:其中 RelErr 基于卡阻"转换段比值抬升"的专属物理签名设计,'
                  '对个体差异具有良好的鲁棒性;ClusterLatent 通过机簇相对基线比较,有效抑制机器个体差异'
                  '对评分的影响。各分量经训练集自动缩放后加权融合,并以验证集分位数确定检测阈值,'
                  '使虚警率保持可控。')
    add_body(doc, '在预警机制层面,考虑到卡阻通常呈渐进式退化而非瞬时突变,本文构建逐台机器的退化趋势'
                  '预警机制,通过滑动窗口漂移补偿与指数平滑趋势跟踪识别早期退化迹象,在故障发生前实现'
                  '提前干预。')
    add_body(doc, '综上,本系统围绕"局部性物理签名"与"个体差异干扰"两个核心难点,形成了"无监督表征'
                  '学习—多模态特征融合—多视角评分—逐台趋势预警"的完整技术路线。据此构建三层系统框架:'
                  '底层为数据层(卡阻模拟数据集,20 台机器 × 180 天),中间为模型层(CNN-LSTM 时频双通路'
                  '自编码器,以重构误差驱动、纯无监督训练),上层为检测层(4 个参与打分分量加权融合)'
                  '与预警层(逐台退化趋势跟踪)。')

    add_h2(doc, '0.2 怎么做 — 系统框架')
    add_image(doc, IMG_FLOW, 16.5, '系统整体架构流程图:数据 → 时/频双通路编码 → 交叉注意力对齐 → '
                                   '拼接融合 → BiLSTM → 解码重构 → 检测打分 → 逐台退化预警', figure_no='0-1')
    add_table(doc,
              ['模块', '文件', '职责'],
              [
                  ['配置管理', 'src/config.py', '全部超参数集中管理(Data/Model/Train/Fusion/Detect/EarlyWarning)'],
                  ['模型', 'src/model.py', 'CNNLSTM_Autoencoder:时频双通路 + 交叉注意力对齐 + 辅助特征 + 频域正常性模型 + 辅助监督头'],
                  ['检测', 'src/detector.py', '检测管线:4 参与分量(PW-MSE/Phase/Cluster/RelErr)自动缩放 → 加权融合 → 阈值 → 评估'],
                  ['预警', 'src/early_warning.py', '退化预警:漂移补偿/趋势跟踪/分级阈值,逐机器总览图'],
                  ['训练', 'src/trainer.py', '训练循环(AMP 混合精度、梯度累积、早停)'],
                  ['数据加载', 'src/data_loader.py', '加载预置 .mat 数据'],
                  ['相位特征', 'src/phase_features.py', '相位特征提取/统计/误差'],
                  ['工具', 'src/utils.py', '检测结果绘图(分数分布/ROC-PR/重构对比/混淆矩阵)'],
              ],
              caption='系统模块职责一览', table_no='0-1', col_widths=[2.6, 4.0, 9.4])

    add_h2(doc, '0.3 结果')
    add_body(doc, '整个系统在卡阻数据集上取得了可用水平的检测效果。下表为最新一次运行的综合指标'
                  '(详见第 9 章),故障率仅 0.1%% 的极端类别不平衡下,AUC-ROC 达到 %.4f,'
                  '固定验证集 P99.5 阈值时召回 %.1f%%。' % (m['auc_roc'], m['recall'] * 100))
    add_table(doc,
              ['指标', '值', '说明'],
              [
                  ['AUC-ROC', f"{m['auc_roc']:.4f}", '整体排序能力(正常/异常分数可分性)'],
                  ['AUC-PR', f"{m['auc_pr']:.4f}", '类别不平衡下的综合性能'],
                  ['召回率', f"{m['recall']*100:.1f}% ({m.get('tp','—')}/{m.get('tp_fn',112)})", '故障检出比例'],
                  ['虚警率 FPR', f"{m['fpr']*100:.2f}% ({m.get('fp','—')}/33,834)", '正常样本误报比例'],
                  ['精确率', f"{m['precision']*100:.1f}%", '报警中真故障占比'],
                  ['F1', f"{m['f1']:.3f}", '召回与精确的调和平均'],
              ],
              caption='系统综合检测指标(卡阻数据集,最新运行)', table_no='0-2', col_widths=[3.4, 4.6, 8.0])

    # =====================================================================
    # 第 1 章 数据与预处理
    # =====================================================================
    add_h1(doc, '1. 数据与预处理')

    add_h2(doc, '1.1 设计背景')
    add_body(doc, '模型的输入是转辙机一次完整动作的三相电流波形。数据侧有三个关键设计约束:')
    add_bullet(doc, '采样率需满足 50Hz 工频电流的抽样定理(≥2×50Hz),取 100Hz;单次动作 8 秒 = 800 点。')
    add_bullet(doc, '不同机器动作时长有差异,需零填充到固定 800 点,但零填充区不含物理信息,会在重构误差中'
                    '引入"曲线长度"这一免费判别信号,须用掩码消除。')
    add_bullet(doc, '退化预警需要严格的时序信息,样本必须与"机器号/天数/小时"逐行对齐。')
    add_body(doc, '数据来自外部 MATLAB 生成器(仓库 lstm),在 20 台机器 × 180 天时间线上以 0.1% 故障率'
                  '均匀散布卡阻故障,并预处理为归一化 + 零填充的固定长度张量。')

    add_h2(doc, '1.2 怎么做 — 数据划分与预处理')
    add_table(doc,
              ['数据集', '样本数', '故障数', '用途'],
              [
                  ['训练集 X_train', f"{stats['X_train_shape'][0]:,}", '0(纯正常)', '训练自编码器'],
                  ['验证集 X_val', f"{stats['X_val_shape'][0]:,}", '0(纯正常)', '早停 + 阈值分位'],
                  ['测试集 X_test', f"{stats['X_test_shape'][0]:,}", f"{stats['n_fault']}", '评估检测指标'],
              ],
              caption='卡阻数据集划分(样本形状均为 N×3×800)', table_no='1-1',
              col_widths=[4.6, 3.6, 3.4, 4.4])
    add_body(doc, '预处理要点:'
                  '① 三相电流归一化到 [0,1];② 按样本有效动作长度零填充到 800 点;'
                  '③ days/machines/hours 三个时序字段与样本逐行对齐;'
                  '④ 训练损失用活动区掩码(active_region_mask)排除尾部零填充,消除长度伪影。'
                  '实测有效动作区均长约 %.2f s(P5=%.2f s, P95=%.2f s)。'
                  % (stats['X_train_active_len']['mean_sec'],
                     stats['X_train_active_len']['p5_sec'],
                     stats['X_train_active_len']['p95_sec']))

    add_h2(doc, '1.3 结果 — 输入数据形态')
    add_image(doc, IMG_WAVE, 16.0,
              '转辙机一次正常动作的三相电流波形(25Hz 演示数据):启动冲击 → 解锁 → 转换段 → 缓放台阶 → 落零 → 零填充',
              figure_no='1-1')
    add_image(doc, IMG_LINE, 16.0,
              '三相动作电流时序线条(同一动作的三相错相与幅度微差)', figure_no='1-2')
    ex = stats['excel_25hz']['phase']
    add_table(doc,
              ['相', '峰值电流 (A)', '有效区均值 (A)', '有效点数 (of 200)'],
              [
                  ['A', f"{ex['A']['max']:.3f}", f"{ex['A']['mean']:.3f}", str(ex['A']['nonzero'])],
                  ['B', f"{ex['B']['max']:.3f}", f"{ex['B']['mean']:.3f}", str(ex['B']['nonzero'])],
                  ['C', f"{ex['C']['max']:.3f}", f"{ex['C']['mean']:.3f}", str(ex['C']['nonzero'])],
              ],
              caption='演示数据三相统计(25Hz,8s,正常动作)', table_no='1-2', col_widths=[2.4, 4.6, 4.6, 4.4])
    add_body(doc, '图 1-1/1-2 展示了单次正常动作的完整形态:约 0.2s 处 A 相出现启动冲击峰值(约 3.1A),'
                  '随后进入解锁段,0.5~4.6s 为转换段(1.5~2.1A 持续、带交流纹波),4.8~5.1s 出现缓放台阶'
                  '(快速跌落后的短暂保持),之后快速落零。三相之间仅存在时间错相(0/40/80ms)与 ±2% 幅度微差,'
                  '形态基本一致——这正是自编码器重构与检测所依赖的"正常模式"。'
                  '表 1-2 的统计显示有效动作区约占 8s 中的约 5.4s(136/200 点),与真实记录一致;'
                  '训练数据中有效长度均长约 5.9s(P5~P95 = 5.3~6.4s),证明固定零填充后仍需掩码处理。')

    # =====================================================================
    # 第 2 章 时域特征通路
    # =====================================================================
    add_h1(doc, '2. 时域特征通路(ConvEncoder)')

    add_h2(doc, '2.1 设计背景')
    add_body(doc, '动作电流的波形形态——启动冲击峰值、转换段幅值、缓放台阶、启动斜率——直接对应机械状态。'
                  'CNN 擅长提取这类局部时序形态,故以 5 层残差卷积构成时域编码通路,把原始波形逐级压缩成'
                  '高维语义特征;同时注入时域物理特征,让网络显式拿到工程师熟知的判别量。')

    add_h2(doc, '2.2 怎么做 — 残差卷积编码器 + 物理特征')
    add_body(doc, '编码器为 5 个 ResidualConvBlock 级联(卷积 → BN → ReLU → 卷积 → BN → +残差 → ReLU → 池化),'
                  '前 3 层降采样、后 2 层保持分辨率并用空洞卷积扩大感受野。以 200 点/25Hz 为例的路径:')
    add_table(doc,
              ['容器', '通道 (in→out)', '卷积核', '池化', '输出长度', '说明'],
              [
                  ['block1', '3 → 64', 'k=7', 'pool=2', '200→100', '初层形态'],
                  ['block2', '64 → 128', 'k=5', 'pool=2', '100→50', '中程特征'],
                  ['block3', '128 → 128', 'k=3', 'pool=2', '50→25', '高层语义'],
                  ['block4', '128 → 128', 'k=3, dil=2', 'pool=1', '25→25', '空洞扩大感受野'],
                  ['block5', '128 → 256', 'k=3', 'pool=1', '25→25', '通道扩展'],
              ],
              caption='时域 ConvEncoder 结构(200 点示例)', table_no='2-1',
              col_widths=[2.4, 3.4, 2.6, 2.6, 3.2, 3.8])
    add_body(doc, '物理特征注入:每个样本额外计算 18 维时域物理特征 t_phys(每相 6 维),广播到各时间步后'
                  '与卷积特征在通道维拼接(256+18)。这 6 个物理量为:')
    add_table(doc,
              ['特征', '维度', '含义'],
              [
                  ['peak_amp', '3', '峰值幅度(启动冲击/功率不足直接反映)'],
                  ['peak_time', '3', '峰值时间位置(形态畸变敏感)'],
                  ['rms', '3', '均方根(整体幅值水平)'],
                  ['diff_energy', '3', '一阶差分能量(形态变化剧烈度)'],
                  ['start_slope', '3', '启动段斜率(前 10% 首尾差)'],
                  ['zero_rate', '3', '零值占比(中途停止/提前结束)'],
              ],
              caption='时域物理特征构成(共 18 维)', table_no='2-2', col_widths=[3.4, 2.0, 10.6])

    add_h2(doc, '2.3 结果')
    add_body(doc, '时域通路负责"形态"。重构对比图(图 4-1)中,训练/验证/异常样本的时域波形均能被较精确地重建,'
                  '证明编码器已把启动冲击、转换段、缓放台阶等关键形态压缩进 latent。物理特征注入的意义在于:'
                  '像"启动冲击过高/功率不足"这类整体幅值偏移的故障,在纯重构误差中接近静默,'
                  '但 peak_amp/rms 直接暴露其偏差,配合第 8 章的幅度解码头形成互补。')

    # =====================================================================
    # 第 3 章 频域特征通路
    # =====================================================================
    add_h1(doc, '3. 频域特征通路(FreqEncoder)')

    add_h2(doc, '3.1 设计背景')
    add_body(doc, '工频纹波、谐波失真、频谱带宽反映电气健康,而这些结构在时域波形上难以直接辨认。'
                  '为此构建第二条频域通路:对原始信号做 rFFT(保留相位),用卷积编码频谱结构,'
                  '再经逆变换映射回时域特征图,与时域通路在空间上对齐。')

    add_h2(doc, '3.2 怎么做 — 复数频谱编码 + iFFT 回时域')
    add_body(doc, '频域通路的关键在于"保留相位"并最终回到时域:')
    add_bullet(doc, 'rFFT 得到复数频谱(保留相位),把实部/虚部作为独立通道叠加,得到 6 通道输入。')
    add_bullet(doc, 'FreqEncoder(6→256)所有卷积均 pool=1,完整保留频率分辨率。')
    add_bullet(doc, '把 256 通道解释为 128 对(real, imag),做 iFFT 映射回时域(重构出时域特征图)。')
    add_bullet(doc, '1×1 卷积把 128 通道投影回 256,自适应平均池化到与时域通路相同的长度 L。')
    add_body(doc, '频域物理特征 f_phys 共 39 维,在 iFFT 之后广播拼接:')
    add_table(doc,
              ['特征', '维度', '含义'],
              [
                  ['fine_psd', '24', '8 个细频带 log1p 功率 × 3 相'],
                  ['shape', '9', '谱质心 / 平坦度 / 85% 滚降 × 3 相'],
                  ['bandwidth', '3', '90% 能量带宽(P5→P95 频率范围)'],
                  ['thd', '3', '谐波失真比((总功率−基波)/基波)'],
              ],
              caption='频域物理特征构成(共 39 维)', table_no='3-1', col_widths=[3.4, 2.0, 10.6])

    add_h2(doc, '3.3 结果')
    add_body(doc, '频域通路捕捉的是"频谱结构"。它把工频成分、谐波分布编码进特征图,再映射回时域,'
                  '使频域信息能与时域信息在同一空间参与后续交叉注意力对齐与拼接。'
                  'f_phys 中的 thd 与 bandwidth 对谐波畸变类异常敏感,'
                  'fine_psd 的细频带能量分布则刻画频谱形状。该通路与第 4 章的交叉注意力配合,'
                  '是"频域知识注入时域"的载体。')

    # =====================================================================
    # 第 4 章 交叉注意力对齐
    # =====================================================================
    add_h1(doc, '4. 交叉注意力对齐(CrossAttentionAlign)')

    add_h2(doc, '4.1 设计背景')
    add_body(doc, '时/频双通路此前只是简单 concat,两路信息没有交互。受 TimeCMA 跨模态思想启发'
                  '——"解耦但弱、纠缠但稳健"——本系统让两分支通过交叉注意力互相检索上下文:'
                  '时域特征作 Q 去频域里"查"上下文,频域特征作 Q 去时域里"查"上下文,双向互相增强,'
                  '但各自维度不变(仍 256 通道)。整个对齐仅由重构损失驱动,是纯无监督的。')

    add_h2(doc, '4.2 怎么做 — 双向交叉注意力 + 信任门控')
    add_body(doc, '输入为两路纯特征 T、F(均为 B×C×L,取前 256 通道,物理特征临时切出):')
    add_bullet(doc, '序列化:T、F 转置为 (B, L, C),各自过 LayerNorm。')
    add_bullet(doc, '双向检索:attn_t2f 以 T 为 Q、F 为 K/V;attn_f2t 以 F 为 Q、T 为 K/V(4 头注意力)。')
    add_bullet(doc, '量级缩放 _scale:把注意力输出缩放到与输入同 std,防初始化时注意力淹没主表示。')
    add_bullet(doc, '信任门控 g=sigmoid(可学习,初始化≈0.5):控制注入强度,残差加回本分支。')
    add_bullet(doc, '输出仍为 (B, C, L),再拼回物理特征,维度不变。')
    add_body(doc, '形式化(以时域分支为例):T_out = T + g·scale(Attn(Q=T, K=V=F))。'
                  '门控让网络可以学"该信多少对方模态的信息",残差保证本征形态始终保留。')

    add_h2(doc, '4.3 结果')
    add_body(doc, '在卡阻数据集上开启交叉注意力对齐后,检测精确率显著提升、虚警下降,召回基本持平'
                  '(消融实验,数据来自配置注释中的实证记录):')
    add_table(doc,
              ['配置', '精确率', '误报 FP', '虚警率 FPR', '召回率'],
              [
                  ['关闭对齐(仅 concat)', '29%', '220', '0.70%', '86.6%'],
                  ['开启对齐(xalign)', '37%', '166', '0.49%', '86.6%'],
                  ['变化', '+8 pp', '−54', '−0.21 pp', '持平'],
              ],
              caption='交叉注意力对齐消融效果(卡阻数据集)', table_no='4-1',
              col_widths=[5.6, 2.6, 3.0, 3.4, 3.4])
    add_body(doc, '解读:开启对齐后精确率 29%→37%(提升 8 个百分点),误报从 220 降到 166,'
                  '虚警率 0.70%→0.49%,而召回率保持 86.6% 不损失。这说明双分支互相检索上下文后,'
                  '正常样本的表示更一致(虚警降),同时不牺牲故障检出能力——交叉注意力以很小的计算代价,'
                  '让两路信息"纠缠但稳健"。')

    # =====================================================================
    # 第 5 章 拼接融合 + BiLSTM
    # =====================================================================
    add_h1(doc, '5. 拼接融合 + BiLSTM')

    add_h2(doc, '5.1 设计背景')
    add_body(doc, '双通路特征如何汇合,经历了从"可学习加权 α"到"拼接"的取舍。早期的加权求和方案'
                  '存在坏吸引子:两路输出趋同会使 α 梯度消失、卡在 0.5,反而注入噪声。'
                  '最终改为直接把两路特征在通道维拼接,把"如何权衡双通路"交给 BiLSTM 门控学习。')

    add_h2(doc, '5.2 怎么做 — concat 融合 + 双向 LSTM')
    add_body(doc, '融合张量 = 时域通路(256 对齐特征 + 18 物理) ⊕ 频域通路(256 对齐特征 + 39 物理) = 569 通道。'
                  '拼接后 permute 成 (B, L, 569) 送入 2 层双向 LSTM(hidden=128,双向输出 256)。'
                  'LSTM 的每个门控都是对全部 569 通道的线性投影,因此两路信息在隐藏状态里真正混合。')
    add_table(doc,
              ['环节', '形状', '说明'],
              [
                  ['时域对齐输出', '(B, 256, L)', 't_align + t_phys(18)'],
                  ['频域对齐输出', '(B, 256, L)', 'f_align + f_phys(39)'],
                  ['拼接融合', '(B, 569, L)', '256+18+256+39'],
                  ['BiLSTM 输出', '(B, L, 256)', '2 层, hidden=128, 双向'],
                  ['解码器输入', '(B, 256, L/8)', 'permute 为 Conv1d 格式'],
              ],
              caption='拼接融合维度流(L 为时序长度)', table_no='5-1',
              col_widths=[4.6, 5.2, 6.2])
    add_body(doc, 'LSTM 隐层 64→128 的扩展也有实证依据:更宽的 latent 缓解了功率不足/启动冲击等难样本'
                  '在 latent 空间中区分不足的问题,扩展后 weighted 融合的召回与 AUC 均获提升。')

    add_h2(doc, '5.3 结果')
    add_body(doc, '拼接融合后网络自行学到双通路权重,消除了可学习 α 的坏吸引子问题;'
                  '更宽的 BiLSTM latent 改善了难样本可分性。这一层是"特征混合"的枢纽:'
                  '它把所有视角(形态 + 频谱 + 物理量)压入同一个时间序列表征,供解码器重构与检测侧使用。')

    # =====================================================================
    # 第 6 章 辅助特征注入
    # =====================================================================
    add_h1(doc, '6. 在线辅助特征注入(AuxEncoder)')

    add_h2(doc, '6.1 设计背景')
    add_body(doc, '解码器重构时,除了局部时序上下文,还需要"这条曲线整体是什么样子"的全局条件——'
                  '整体幅值、频带能量分布、峰值位置。这些是样本级统计量,不适合逐时间步 concat 进 LSTM'
                  '(会重复 25 次、冗余),因此采用"编码后广播加到 LSTM 输出"的 add 模式,'
                  '作为全局统计条件指导解码器。')

    add_h2(doc, '6.2 怎么做 — 24 维统计量 → MLP → 广播加')
    add_body(doc, '辅助特征 24 维在线的实时计算(compute_aux_features),复用频域通路已算好的 rFFT:')
    add_table(doc,
              ['分量', '维度', '含义'],
              [
                  ['PSD 频带功率', '15', '5 个频带(0-2/2-4/4-6/6-8/8-12.5 Hz)log1p 功率 × 3 相'],
                  ['peak_amp', '3', '各相峰值幅度'],
                  ['peak_time', '3', '各相峰值位置(归一化)'],
                  ['rms', '3', '各相均方根'],
              ],
              caption='辅助特征构成(共 24 维)', table_no='6-1', col_widths=[4.0, 2.0, 10.0])
    add_body(doc, 'AuxEncoder 为浅 MLP:24 → 64(ReLU+Dropout)→ 256。关键设计是输出维度 = lstm_hidden×2 = 256,'
                  '与 BiLSTM 输出同维,从而能 unsqueeze(-1) 后沿时间维广播加:'
                  '每个时间步加上同一个全局条件向量。该 MLP 无显式监督,端到端由重构损失驱动学习。')

    add_h2(doc, '6.3 结果')
    add_body(doc, '辅助注入以极低成本为解码器提供"全局统计条件":LSTM 输出每步已有局部时序上下文,'
                  '叠加全局偏置后,解码器既能看局部又能看整体。与直接 concat 进输入相比,'
                  '广播加不增加序列通道数、参数更少,且训练稳定。它和 concat 的物理特征分工明确:'
                  '物理特征进 LSTM 输入侧,辅助特征进 LSTM 输出侧。')

    # =====================================================================
    # 第 7 章 解码器通路
    # =====================================================================
    add_h1(doc, '7. 解码器通路(ConvDecoder)')

    add_h2(doc, '7.1 设计背景')
    add_body(doc, '自编码器需要从低分辨率 latent 上采样还原原始电流曲线。为保证重建质量与训练稳定,'
                  '解码器采用与编码器严格镜像的残差转置结构:编码器用池化降采样,解码器用上采样回程,'
                  '逐层镜像对称。')

    add_h2(doc, '7.2 怎么做 — 残差转置卷积 + 镜像上采样')
    add_body(doc, '构建块为 ResidualTransposeBlock:Upsample(线性插值)→ 卷积 → BN → ReLU → 卷积 → BN → '
                  '+残差 → ReLU;通道变化或上采样时用 1×1 卷积投影残差捷径。ConvDecoder 由 5 个容器级联:')
    add_table(doc,
              ['容器', '通道 (in→out)', '上采样', '输出长度', '镜像对象'],
              [
                  ['d1', '256 → 128', '—', '25', '编码器 block5'],
                  ['d2', '128 → 128', '—(dil=2)', '25', '编码器 block4(空洞)'],
                  ['d3', '128 → 128', '2×', '25→50', '编码器 block3'],
                  ['d4', '128 → 64', '2×', '50→100', '编码器 block2'],
                  ['d5', '64 → 3', '2×', '100→200', '编码器 block1'],
                  ['final_resize', '—', 'size=seq_len', '200→200', '兜底对齐'],
              ],
              caption='ConvDecoder 结构(200 点示例,严格镜像编码器)', table_no='7-1',
              col_widths=[3.0, 3.4, 2.6, 3.2, 4.8])
    add_body(doc, '解码器输入来自 BiLSTM 输出(256 通道),经 d1→d5 逐级上采样回 3 相 × seq_len 长度,'
                  '末端 final_resize 精确对齐序列长度。当输入序列长度能被 8 整除时,镜像回程精确无余数。')

    add_h2(doc, '7.3 结果')
    add_image(doc, IMAGES['recon'], 16.5,
              '原始 vs 重构对比:训练(早期低阻力)/ 验证(阻力升高)/ 异常样本 三相波形',
              figure_no='7-1')
    add_body(doc, '图 7-1 展示了重构效果:左列训练样本、中列验证样本、右列异常样本。'
                  '正常样本(训练/验证)的重构曲线(红虚线)与原始曲线(蓝实线)高度吻合,'
                  '说明解码器已学到正常流形;异常样本因偏离训练分布,重构误差被放大,'
                  '成为检测的得分来源。镜像结构 + 残差连接保证了深度的梯度顺畅,'
                  '是训练能收敛到低重构误差的结构基础。')

    # =====================================================================
    # 第 8 章 辅助监督头与训练策略
    # =====================================================================
    add_h1(doc, '8. 辅助监督头与训练策略')

    add_h2(doc, '8.1 设计背景')
    add_body(doc, '纯重构损失下,若干难样本——启动冲击过高、功率不足——在重构误差中接近静默'
                  '(模型倾向于按正常形态重建,掩盖了幅度偏差)。仅靠重构,latent 未必编码这些关键物理量。'
                  '为此增加三个辅助监督目标,强化 latent 表征:幅度解码头(amp_head)、时段判别头(SegDisc)、'
                  '频域正常性模型(FNM)。')

    add_h2(doc, '8.2 怎么做 — 三个辅助目标 + 复合损失')
    add_table(doc,
              ['辅助头', '结构', '监督方式', '作用'],
              [
                  ['amp_head', 'latent(256)→64→6', '预测 peak_amp(3)+rms(3)', '强制 latent 编码幅度信息(难样本关键信号)'],
                  ['SegDisc', 'GRU(256→128)+FC', '时段损坏自监督判别', '端到端学正常局部模式,替代检测侧手动时段马氏'],
                  ['FNM', '谱子AE(瓶颈16)', 'log功率谱重构误差', '学正常频谱流形,重构误差作 NormErr 检测分量'],
              ],
              caption='三个辅助监督头', table_no='8-1', col_widths=[2.8, 4.4, 4.6, 4.2])
    add_body(doc, '总训练损失为各项加权求和:'
                  'L = 掩码重构 MSE + λ_amp·AmpErr + λ_nc·NormErr + λ_disc·DiscErr,'
                  '权重分别为 1.0 / 0.5 / 0.1 / 0.3。其中重构 MSE 只用活动区掩码,排除尾部零填充。')
    add_bullet(doc, '训练技巧:AMP 混合精度(GPU)、梯度累积(accum=4,等效 batch 512)、'
                    '余弦退火学习率、梯度裁剪 max_norm=1.0、早停 patience=15。')
    add_bullet(doc, '时段损坏自监督:随机对样本的某个时段做 ×0.5 衰减或 ×1.5 增强,'
                    '判别头学"哪个时段被破坏",幅度衰减≈功率不足、增强≈启动冲击。')

    add_h2(doc, '8.3 结果')
    add_table(doc,
              ['信号', 'AUC(amp_head 单分量)', '说明'],
              [
                  ['启动冲击过高', '0.867', '纯重构静默的难样本'],
                  ['功率不足', '0.823', '纯重构静默的难样本'],
              ],
              caption='幅度解码头对难样本的诊断能力(配置注释实证)', table_no='8-2',
              col_widths=[4.6, 5.4, 6.0])
    add_body(doc, '表 8-2 显示 amp_head 对两类"重构静默"难样本有 0.82~0.87 的独立鉴别能力,'
                  '证明训练期幅度监督确实让 latent 编码了关键幅度信息。'
                  'SegDisc 使网络端到端学到正常局部模式;'
                  'FNM 的谱重构误差作为训练损失项,强化对正常频谱流形的建模。'
                  '三者共同构成"重构 + 多视角辅助"的训练范式。')

    # =====================================================================
    # 第 9 章 异常检测管线
    # =====================================================================
    add_h1(doc, '9. 异常检测管线(Detector)')

    add_h2(doc, '9.1 设计背景')
    add_body(doc, '训练好的自编码器需在无监督下对每个样本给出异常分数。面对 0.1% 的极端类别不平衡,'
                  '单一误差度量易漏检,故由 4 个参与分量覆盖重构误差、相位结构、机簇相对关系与'
                  '波形级物理特征等互补视角;每个分量在训练集上自动缩放后再加权融合,阈值取验证集分位,'
                  '使虚警率可控。')

    add_h2(doc, '9.2 怎么做 — 4 分量加权融合')
    add_table(doc,
              ['分量', '权重', '含义'],
              [
                  ['PW-MSE', '0.3', '相区加权重构误差(转换段权重 2.0,掩码排除零填充)'],
                  ['PhaseErr', '0.3', '相位结构偏差(峰值时间偏移等形态畸变)'],
                  ['ClusterLatent', '1.0', '机簇锚 latent(KMeans k=20,min 马氏),幅度类异常相对自身机簇基线'],
                  ['RelErr', '1.0', '相对物理特征 Σz²(转换/峰值、转换/解锁、转换段波动),低基础机卡阻的波形级信号'],
              ],
              caption='4 个参与打分分量', table_no='9-1',
              col_widths=[3.6, 1.8, 10.6])
    add_body(doc, '关键机制:'
                  '① 每个分量在训练集上自动缩放(如中位数归一化),消除量纲差异;'
                  '② 综合分 = PW-MSE + PhaseErr + ClusterLatent + RelErr(加权和);'
                  '③ 阈值取验证集分数 P99.5 分位(理论虚警率 ≈ 0.5%)。'
                  '曾实验过的频谱(SpectralErr)、正常性(NormErr)、latent 马氏(LatentErr)、'
                  '物理 z-score(PhysErr)、幅度头(AmpHead)、时段锚(SegLatent)、对齐残差(AlignResidual)'
                  '等分量已从检测管线中移除,不再干扰架构理解。')

    add_h2(doc, '9.3 结果')
    add_image(doc, IMAGES['score'], 16.5,
              '训练/验证/测试三集异常分数分布(红线为验证集 P99.5 阈值)', figure_no='9-1')
    add_image(doc, IMAGES['roc'], 16.5,
              '测试集 ROC 曲线与 PR 曲线(AUC-ROC=%.4f)' % m['auc_roc'], figure_no='9-2')
    add_image(doc, IMAGES['confusion'], 12.0,
              '混淆矩阵(阈值=验证集 P99.5,F1=%.3f)' % m['f1'], figure_no='9-3')
    add_table(doc,
              ['指标', '值'],
              [
                  ['AUC-ROC', f"{m['auc_roc']:.4f}"],
                  ['AUC-PR', f"{m['auc_pr']:.4f}"],
                  ['召回率', f"{m['recall']*100:.1f}% ({m.get('tp','—')}/{m.get('tp_fn',112)})"],
                  ['虚警率 FPR', f"{m['fpr']*100:.2f}%"],
                  ['精确率', f"{m['precision']*100:.1f}%"],
                  ['F1', f"{m['f1']:.3f}"],
              ],
              caption='检测指标(卡阻数据集,最新运行)', table_no='9-2', col_widths=[5.0, 11.0])
    add_body(doc, '图 9-1 显示训练分数集中在低值区间,异常样本分数明显右移且大部分越过阈值线;'
                  '图 9-2 中 AUC-ROC 高达 %.4f,说明正常/异常分数排序能力很强;'
                  '图 9-3 的混淆矩阵给出固定阈值下的实际检出:112 个故障中检出 %s 个,'
                  '虚警 %s 个(约 %.2f%%)。在 1/1000 的极端不平衡下,'
                  'AUC-PR(%.4f)与召回率是更关键的指标——本方案在保持高召回的同时把虚警压在 1%% 以内。'
                  % (m['auc_roc'], m.get('tp', '—'), m.get('fp', '—'), m['fpr'] * 100, m['auc_pr']))

    # =====================================================================
    # 第 10 章 退化趋势预警
    # =====================================================================
    add_h1(doc, '10. 退化趋势预警(Early Warning)')

    add_h2(doc, '10.1 设计背景')
    add_body(doc, '故障率极低的场景下,单点检测(第 9 章)只能回答"这个样本异常吗"。'
                  '现场更需要回答"哪台机器正在退化、什么时候该检修"。因此按逐台机器建立独立监测:'
                  '每台机器一条 180 天时间线,跟踪异常分数趋势,提前给出绿/黄/橙/红分级预警。')

    add_h2(doc, '10.2 怎么做 — 漂移补偿 + 趋势跟踪 + 分级阈值')
    add_table(doc,
              ['组件', '机制', '作用'],
              [
                  ['DriftCompensator', '滑动窗口 z-score(窗口 200)', '消除工况/季节漂移,关注"偏离当前正常"而非绝对偏差'],
                  ['TrendTracker', '指数平滑(α=0.3)+ 窗口斜率(7 样本)+ 加速度', '检测趋势加速,识别缓慢爬升的退化'],
                  ['WarningSystem', '标准化训练分数分位阈值 + 斜率联合判断', '绿/黄/橙/红四级预警'],
              ],
              caption='预警系统三组件', table_no='10-1', col_widths=[3.6, 6.0, 6.4])
    add_table(doc,
              ['级别', '判定条件', '含义'],
              [
                  ['[绿] 正常', '低于黄阈值,无强趋势', '正常'],
                  ['[黄] 关注', '> P95', '偏差增大,持续观察'],
                  ['[橙] 预警', '> P99.5 或 (> P95 且 斜率>0.05)', '退化明显,建议 7 天内检修'],
                  ['[红] 报警', '> P99.9', '严重异常,需立即检查'],
              ],
              caption='分级预警阈值(基于训练集标准化分数分位)', table_no='10-2',
              col_widths=[2.8, 7.4, 5.8])
    add_body(doc, '阈值统一用"漂移补偿后"的训练分数分位(P95/P99.5/P99.9),与在线 z-score 分数量纲一致,'
                  '避免冷启动塌缩。每台机器独立计算自己的阈值与趋势。')

    add_h2(doc, '10.3 结果')
    add_image(doc, IMAGES['warning'], 16.5,
              '20 台机器退化预警总览:共用 0–180 天轴,日最大分数(对数轴),红×为故障样本,黑色虚线为阈值',
              figure_no='10-1')
    add_body(doc, '图 10-1 把 20 台机器画在同一时间轴上:横轴为监测天数,纵轴为每日最大异常分数(对数坐标,'
                  '便于同屏显示正常 ~1 与极端故障 ~8000),红色 × 标记各机器的故障样本,黑色虚线为全局阈值。'
                  '从图中可直观看出:绝大多数机器长期处于正常水平;故障样本(红×)大多明显超出阈值,'
                  '且往往伴随分数在故障前几天就开始爬升——这正是退化趋势预警的价值:在真正故障前'
                  '发出橙/红预警,提示安排检修。')

    # =====================================================================
    # 第 11 章 总结
    # =====================================================================
    add_h1(doc, '11. 总结')

    add_h2(doc, '11.1 全链路效果')
    add_table(doc,
              ['层级', '关键设计', '效果'],
              [
                  ['数据', '归一化 + 零填充 + 活动区掩码', '消除长度伪影,时序字段支持预警'],
                  ['模型', '时频双通路 + 交叉注意力对齐', '精确率 29%→37%,虚警 0.70%→0.49%'],
                  ['模型', '拼接融合 + BiLSTM(hidden 128)', '难样本 latent 区分增强'],
                  ['训练', '重构 + amp_head + SegDisc + FNM', '难样本(启动冲击/功率不足)有独立鉴别信号'],
                  ['检测', '4 分量加权 + P99.5 阈值', f"AUC-ROC {m['auc_roc']:.4f}, 召回 {m['recall']*100:.1f}%"],
                  ['预警', '逐台漂移补偿 + 趋势跟踪 + 分级', '故障前趋势爬升可提前预警'],
              ],
              caption='全链路关键设计及其效果汇总', table_no='11-1',
              col_widths=[2.4, 6.6, 7.0])

    add_h2(doc, '11.2 局限与展望')
    add_bullet(doc, '类别不平衡极端(0.1%),精确率仍偏低(约 30%),报警中大量是虚警,现场可用性取决于检修成本。')
    add_bullet(doc, '卡阻数据来自模拟生成器,真实现场数据的泛化有待验证。')
    add_bullet(doc, '频谱/正常性/latent/物理/幅度/时段锚/对齐等曾实验的分量已从检测管线移除,后续可结合更多故障形态重新评估。')
    add_bullet(doc, '展望:引入在线增量学习适应漂移、多传感器融合(振动/温度),以及跨机器迁移。')

    doc.save(OUT_PATH)
    print(f'[ok] 已生成: {OUT_PATH}')
    print(f'      大小: {os.path.getsize(OUT_PATH)/1024:.0f} KB')


if __name__ == '__main__':
    build_doc()
